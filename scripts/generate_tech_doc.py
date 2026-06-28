"""生成 NFSW_Detector 项目技术文档（docx）。

使用 python-docx 生成完整技术文档，涵盖项目概述、系统架构、SVLA 算法设计、
训练与评估、推理管线、API 设计、部署配置、技术创新点等内容。

用法：
    python scripts/generate_tech_doc.py
输出：
    docs/NFSW_Detector_技术文档.docx
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ==================== 样式辅助 ====================

def set_cell_shading(cell, color_hex):
    """设置单元格底纹颜色。"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_code_block(doc, code_text):
    """添加代码块（等宽字体 + 浅灰底纹）。"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2e, 0x2e, 0x2e)
    # 段落底纹
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F5F5F5')
    p_pr.append(shd)


def add_formula(doc, formula_text):
    """添加公式段落（等宽字体居中）。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(formula_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)


def add_table_from_rows(doc, headers, rows, col_widths=None):
    """从表头和行数据添加表格。"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, '4472C4')
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(10)
    # 数据行
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    # 列宽
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table


def add_toc(doc):
    """插入目录域代码（用户首次打开按 F9 更新）。"""
    p = doc.add_paragraph()
    run = p.add_run()
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_char_separate = OxmlElement('w:fldChar')
    fld_char_separate.set(qn('w:fldCharType'), 'separate')
    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)


# ==================== 文档生成 ====================

def build_document():
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # ==================== 封面 ====================
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('NFSW Detector')
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('多模态有害内容审查与预警系统')
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('基于 Shot-Conditioned Vision-Language Adaptation (SVLA)')
    run.font.size = Pt(13)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    for _ in range(8):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('技术文档 v1.0\n2026-06-28')
    run.font.size = Pt(12)

    doc.add_page_break()

    # ==================== 目录 ====================
    doc.add_heading('目录', level=1)
    add_toc(doc)
    doc.add_page_break()

    # ==================== 第一章 项目概述 ====================
    doc.add_heading('第一章 项目概述', level=1)

    doc.add_heading('1.1 项目背景', level=2)
    doc.add_paragraph(
        '随着短视频社交平台的快速发展，海量用户生成内容（UGC）中夹杂的吸烟、血腥、暴力、'
        '辱骂、色情、金钱诈骗、政治敏感等有害内容对平台生态与未成年人保护构成严峻挑战。'
        '传统人工审核成本高、效率低，单一模态检测易漏检误判。本项目基于多模态视觉-语言'
        '预训练模型 CLIP 与 Shot-Conditioned Vision-Language Adaptation (SVLA) 算法，'
        '构建端到端的视频有害内容审查与预警系统，实现对短视频内容的自动检测、类别判定、'
        '风险分级与结构化预警报告生成。'
    )

    doc.add_heading('1.2 项目目标', level=2)
    goals = [
        '7 类有害内容检测：吸烟、血腥、暴力、辱骂、色情、金钱诈骗、政治敏感',
        '二分类异常分数：视频级异常概率 [0,1]，基于段级 sigmoid 最大值',
        '多类别条件概率：anomaly_score × P(cat|anomaly)，保证类别分数总和 ≤ anomaly_score',
        '4 级预警等级：SAFE / LOW / MEDIUM / HIGH，对应不同处置建议',
        '关键帧抽取：每个有害段抽取峰值帧保存为 jpg，base64 内联到 HTML 报告',
        '结构化预警报告：JSON / HTML / CSV 三种导出格式，含报告 ID (UUID4)',
        'RESTful API：FastAPI 实现，支持视频与图片检测，OpenAPI 文档自动生成',
    ]
    for g in goals:
        doc.add_paragraph(g, style='List Bullet')

    doc.add_heading('1.3 有害内容类别定义', level=2)
    add_table_from_rows(doc,
        ['ID', '英文', '中文', '描述'],
        [
            ['0', 'Smoke', '吸烟', '检测到疑似吸烟行为，可能违反平台内容规范'],
            ['1', 'Blood', '血腥', '检测到血腥画面，可能引起观众不适'],
            ['2', 'Violent', '暴力', '检测到暴力行为，违反平台社区准则'],
            ['3', 'Abusive', '辱骂', '检测到辱骂行为，可能构成言语骚扰'],
            ['4', 'Sexy', '色情', '检测到色情内容，违反平台内容政策'],
            ['5', 'Money', '金钱诈骗', '检测到疑似金钱诈骗内容，存在欺诈风险'],
            ['6', 'Policy', '政治敏感', '检测到政治敏感内容，可能违反相关规定'],
        ],
        col_widths=[0.5, 1.2, 1.0, 3.5]
    )

    doc.add_heading('1.4 性能指标目标', level=2)
    add_table_from_rows(doc,
        ['指标', '目标值', '说明'],
        [
            ['视频级 AUC', '≥ 0.85', '异常二分类 ROC 曲线下面积'],
            ['视频级 AP', '≥ 0.80', '异常二分类 PR 曲线下面积'],
            ['推理速度', '≤ 5s / min 视频', '含解码、特征提取、SVLA 前向、报告生成'],
            ['类别覆盖', '7 类', 'Smoke / Blood / Violent / Abusive / Sexy / Money / Policy'],
        ],
        col_widths=[1.5, 1.5, 3.5]
    )

    doc.add_page_break()

    # ==================== 第二章 系统架构 ====================
    doc.add_heading('第二章 系统架构', level=1)

    doc.add_heading('2.1 顶层模块依赖', level=2)
    doc.add_paragraph('项目按功能分为 8 层，模块依赖关系如下：')
    add_code_block(doc,
        'main.py (CLI 入口)\n'
        '  ├─ engine/train.py      → 训练引擎\n'
        '  │    └─ data/dataset.py  → SVADataset + DataLoaders\n'
        '  ├─ engine/evaluate.py    → 评估引擎\n'
        '  ├─ pipeline/inference.py → NSFWDetector 端到端推理\n'
        '  │    ├─ pipeline/preprocess.py       → VideoPreprocessor\n'
        '  │    ├─ pipeline/feature_extractor.py → CLIPFeatureExtractor\n'
        '  │    ├─ pipeline/calibration.py       → ScoreCalibrator\n'
        '  │    ├─ models/svla.py                → SVLA 模型\n'
        '  │    └─ pipeline/alert.py             → AlertGenerator\n'
        '  ├─ demo/app.py          → Gradio Demo\n'
        '  ├─ api/app.py           → FastAPI 服务\n'
        '  └─ utils/tools.py       → 公共工具'
    )

    doc.add_heading('2.2 模块职责表', level=2)
    add_table_from_rows(doc,
        ['层级', '模块', '核心类/函数', '职责'],
        [
            ['入口', 'main.py', 'CLI 子命令分发', 'train/evaluate/detect/demo/serve/export'],
            ['数据', 'data/dataset.py', 'SVADataset', 'SVA 数据集加载，双 loader 同步'],
            ['模型', 'models/svla.py', 'SVLA, CFATextAdapter', '主模型 + 跨模态适配器'],
            ['模型', 'models/layers.py', 'GraphConvolution', '图卷积层 + 距离邻接'],
            ['引擎', 'engine/train.py', 'Trainer', '训练循环 + MIL 损失'],
            ['引擎', 'engine/evaluate.py', 'Evaluator', 'AUC/AP 评估'],
            ['管线', 'pipeline/inference.py', 'NSFWDetector', '端到端推理 + 4 项增强'],
            ['管线', 'pipeline/preprocess.py', 'VideoPreprocessor', '视频解码/采样/抽帧'],
            ['管线', 'pipeline/feature_extractor.py', 'CLIPFeatureExtractor', 'CLIP 视觉/文本特征'],
            ['管线', 'pipeline/calibration.py', 'ScoreCalibrator', 'Isotonic Regression 校准'],
            ['管线', 'pipeline/alert.py', 'AlertGenerator', '预警报告生成与导出'],
            ['API', 'api/app.py', 'AppState + 6 路由', 'FastAPI RESTful 服务'],
            ['Demo', 'demo/app.py', 'create_app', 'Gradio Web UI'],
            ['工具', 'utils/tools.py', 'get_prompt_text', 'SVLA 工具函数'],
        ],
        col_widths=[0.8, 2.2, 2.0, 2.5]
    )

    doc.add_heading('2.3 推理数据流', level=2)
    doc.add_paragraph('从视频文件到预警报告的完整推理链路：')
    add_code_block(doc,
        '[视频文件]\n'
        '    │  NSFWDetector.detect(path)\n'
        '    ▼\n'
        '[VideoPreprocessor]   解码、采样、抽帧 → frames [N, H, W, 3]\n'
        '    │\n'
        '[CLIPFeatureExtractor]  视觉特征提取 → features [N, 512]\n'
        '    │\n'
        '[SVLA 模型]  logits1=异常分数, logits2=8类分布\n'
        '    │\n'
        '[推理增强]  校准 + OOD + 质量加权 + 零样本 re-rank\n'
        '    │\n'
        '[DetectionResult]  段级分数 + 类别分数 + 有害时间段 + 关键帧\n'
        '    │\n'
        '[AlertGenerator]  规则化整理、生成摘要与处置建议\n'
        '    │\n'
        '[AlertReport]  最终面向用户展示的预警报告\n'
        '    │\n'
        '[export_json / export_html]  持久化为 JSON/HTML'
    )

    doc.add_heading('2.4 训练数据流', level=2)
    doc.add_paragraph('训练采用 normal + anomaly 双 loader 同步迭代，MIL 损失聚合段级分数到视频级：')
    add_code_block(doc,
        '[normal_loader] ──┐\n'
        '                   ├─→ SVLA.forward(features, text) → logits1, logits2\n'
        '[anomaly_loader] ─┘\n'
        '    │\n'
        '    ▼\n'
        '[π-TopK 软 MIL 池化]  镜头密度自适应 k 值\n'
        '    │\n'
        '    ▼\n'
        '[CLAS2 损失]  二分类异常 (Focal BCE)\n'
        '[CLASM 损失]  多类别分类 (Focal NLL)\n'
        '[TextReg 损失] 文本正则 (类间分离)\n'
        '    │\n'
        '    ▼\n'
        '[总损失]  loss = CLAS2 + α·CLASM + β·TextReg\n'
        '    │\n'
        '    ▼\n'
        '[AdamW + MultiStepLR]  反向传播 + 梯度裁剪 + 早停'
    )

    doc.add_page_break()

    # ==================== 第三章 SVLA 算法设计 ====================
    doc.add_heading('第三章 SVLA 算法设计', level=1)

    doc.add_heading('3.1 整体架构', level=2)
    doc.add_paragraph(
        'SVLA (Shot-Conditioned Vision-Language Adaptation) 是基于 CLIP 的视频异常检测模型，'
        '核心思想是：冻结 CLIP 预训练参数，通过 Shot-Aware 机制与双向 CFA 跨模态适配器'
        '微调适配层，实现视频时序建模与视觉-语言对齐。整体架构包含：'
    )
    arch_points = [
        '视觉编码器：CLIP ViT 提取帧特征 → 帧位置编码 → Temporal Transformer（局部窗口注意力）',
        '镜头条件机制：镜头边界检测 → Shot Transformer → ShotDensityHead（镜头异常密度 π）',
        '双向 CFA 适配：CFATextAdapter（视觉→文本）+ TextCFAdapter（文本←视觉，FiLM）',
        '双分支图卷积：shot-aware 邻接 + 距离邻接，concat 后 MLP 残差',
        '双头输出：logits1（二分类异常 logit）+ logits2（8 类含 normal 的相似度）',
    ]
    for p in arch_points:
        doc.add_paragraph(p, style='List Bullet')

    doc.add_heading('3.2 Shot-Aware 镜头条件机制', level=2)
    doc.add_paragraph(
        '镜头（Shot）是视频的基本叙事单元，同一镜头内的帧语义连贯。SVLA 通过镜头边界检测、'
        '镜头级 Transformer 与镜头密度估计，将镜头信息注入异常检测。'
    )

    doc.add_heading('3.2.1 镜头边界检测', level=3)
    doc.add_paragraph(
        '基于相邻帧 CLIP 特征的余弦相似度检测镜头边界。若相似度低于阈值 shot_sim_thresh '
        '（默认 0.90），则标记为边界。长度小于 shot_min_len（默认 12）的过短镜头会被合并到前一镜头，'
        '避免镜头碎片化。'
    )
    add_code_block(doc,
        'def _detect_shots(self, visual_features, padding_mask, lengths):\n'
        '    # 计算相邻帧余弦相似度\n'
        '    sim = cosine_sim(visual_features[:-1], visual_features[1:])\n'
        '    boundaries = (sim < shot_sim_thresh).nonzero()\n'
        '    # 合并过短镜头\n'
        '    shots = merge_short_shots(boundaries, min_len=shot_min_len)\n'
        '    return shots'
    )

    doc.add_heading('3.2.2 Shot Transformer 与密度估计', level=3)
    doc.add_paragraph(
        '每个镜头内的帧特征取均值得到镜头 token，送入 Shot Transformer 做段内增强。'
        'ShotDensityHead 输出镜头异常密度 π ∈ [pi_floor, 1-ε]，初始化 bias=-1.386 使 '
        'sigmoid 输出 ≈0.2（先验假设多数镜头正常）。π 在训练时与异常标签联合学习，'
        '在推理时作为镜头异常的软先验。'
    )
    add_formula(doc, 'π = sigmoid(ShotDensityHead(shot_token)) ∈ [pi_floor, 1-ε]')

    doc.add_heading('3.3 双向 CFA 跨模态适配器', level=2)
    doc.add_paragraph(
        'CFA (Cross-modal Feature Adaptation) 是 SVLA 的核心创新，实现视觉与语言模态的双向适配。'
        '传统方法仅做单向适配（视觉→文本或文本→视觉），SVLA 同时进行两个方向的适配，'
        '使视觉特征更具语义判别力，文本特征更贴合视觉内容。'
    )

    doc.add_heading('3.3.1 CFATextAdapter（视觉→文本）', level=3)
    doc.add_paragraph(
        'CFATextAdapter 通过动态 Prefix 生成、交叉注意力与门控融合，将文本语义注入视觉特征。'
        '核心流程：'
    )
    cfa_steps = [
        '动态 Prefix 生成：从视觉特征均值经 LayerNorm + 生成网络产生低秩 prefix (prefix_len, prefix_rank)，再投影回 d_vis',
        '静态 + 动态 Prefix 融合：Pk = prefix_k(静态) + Pk_dyn(动态)',
        '交叉注意力：Q=视觉特征，K/V=文本特征 + Prefix，多头注意力',
        '门控融合：F_hat = MLP_bottleneck(F_att)，F_txt_ctx = softmax(V·T^T/τ)·V',
        'F_mod = sigmoid(W_mod(F_txt_ctx))，F_fused = LN(V + β·(F_hat * F_mod))',
    ]
    for s in cfa_steps:
        doc.add_paragraph(s, style='List Number')

    add_formula(doc, 'F_fused = LN(V + β · (MLP_bottleneck(CrossAttn(V, T+Prefix)) ⊙ sigmoid(W_mod(GatedCtx))))')

    doc.add_paragraph(
        '其中 τ (tau) 是门控温度（默认 0.8，clamp 到 [0.25, 4.0]），β (beta) 是融合系数'
        '（默认 0.8，clamp 到 [0,1]），均通过 cfa_tau/cfa_beta 配置。'
    )

    doc.add_heading('3.3.2 TextCFAdapter（文本←视觉，FiLM）', level=3)
    doc.add_paragraph(
        'TextCFAdapter 采用 FiLM (Feature-wise Linear Modulation) 风格，从视觉特征生成'
        '调制系数 (γ, bias) 对文本特征做仿射变换，使文本特征自适应视觉内容：'
    )
    add_formula(doc, 'T_cfa = LN(T_base · (1 + γ) + bias)')
    doc.add_paragraph(
        'γ 和 bias 由视觉特征均值经两层 MLP 生成，初始值使 T_cfa ≈ T_base（训练初期保持文本语义稳定）。'
    )

    doc.add_heading('3.4 π-TopK 软 MIL 池化', level=2)
    doc.add_paragraph(
        '传统 MIL (Multiple Instance Learning) 采用 top-k 池化，k 值固定。SVLA 提出 π-TopK，'
        '根据镜头异常密度 π 自适应调节 k 值——异常密度高的镜头取更多帧（更细致），'
        '正常镜头取更少帧（更高效）。'
    )
    add_formula(doc, 'k_base = max(1, int(L / 16) + 1)')
    add_formula(doc, 'scale = clamp(1 + 2.0 · (π - 0.22), 0.5, 2.0)')
    add_formula(doc, 'k = round(k_base · scale)')
    add_formula(doc, 'pooled_score = topk(scores, k).mean()')

    doc.add_paragraph(
        '其中 L 是镜头长度，π0=0.22 是异常先验（对应 sigmoid(-1.386)），16 是 base_div 参数。'
        '当 π > 0.22 时 scale > 1（异常镜头取更多帧），π < 0.22 时 scale < 1（正常镜头取更少帧）。'
    )

    doc.add_heading('3.5 Shot-Aware 图卷积邻接', level=2)
    doc.add_paragraph(
        'SVLA 使用双分支图卷积融合时序关系。一支用 shot-aware 邻接矩阵，另一支用距离邻接矩阵，'
        '结果 concat 后经 linear + MLP 残差。'
    )

    doc.add_heading('3.5.1 Shot-Aware 邻接矩阵构造', level=3)
    add_formula(doc, 'A_cos = threshold(cos_sim(feat, feat), thr=0.7)')
    add_formula(doc, 'A_shot = block_diag([ones(L_i, L_i) for each shot])')
    add_formula(doc, 'A = A_cos + γ · A_shot + I')
    add_formula(doc, 'A_norm = D^(-1/2) · A · D^(-1/2)')

    doc.add_paragraph(
        '其中 γ 是 shot_gamma（默认 0.05），D 是度矩阵，对称归一化保证数值稳定。'
        'shot-aware 邻接让同镜头帧全连接，强化镜头内信息传播。'
    )

    doc.add_heading('3.5.2 距离邻接矩阵', level=3)
    doc.add_paragraph(
        'DistanceAdj 基于帧间城市距离的指数衰减构造：'
    )
    add_formula(doc, 'A_dist[i,j] = exp(-|i-j| / exp(σ))')
    doc.add_paragraph('σ 可学习，初始值使衰减尺度适中。距离邻接捕获全局时序依赖。')

    doc.add_heading('3.6 文本 prompt 编码', level=2)
    doc.add_paragraph(
        'SVLA 的文本端使用可学习的 text_prompt_embeddings (77 × embed_dim) 作为基础位置嵌入，'
        '按 prompt_prefix/postfix 槽位将真实 CLIP token embedding 覆盖到对应位置，'
        '再调用本地 CLIP 双参数 encode_text(base, text_tokens) 编码。'
    )
    add_code_block(doc,
        'def encode_textprompt(self, text):\n'
        '    tokens = clip.tokenize(text).to(device)\n'
        '    word_embedding = self.clipmodel.encode_token(tokens)\n'
        '    base = self.text_prompt_embeddings(arange(L))\n'
        '    base[prefix_slots] = word_embedding[prefix_slots]\n'
        '    base[postfix_slots] = word_embedding[postfix_slots]\n'
        '    text_features = self.clipmodel.encode_text(base, text_tokens)\n'
        '    return text_features'
    )
    doc.add_paragraph(
        '注意：get_prompt_text() 仅取每类第 1 条 prompt，因此训练后 text_prompt_embeddings '
        '与 prompt-1 耦合。推理增强中的 CLIP 零样本 re-rank 用纯 CLIP 文本编码绕开此偏置。'
    )

    doc.add_heading('3.7 可配置骨干网络', level=2)
    doc.add_paragraph(
        'SVLA 支持通过 YAML 配置切换 CLIP 骨干网络，已验证两种配置：'
    )
    add_table_from_rows(doc,
        ['配置', 'clip_variant', 'embed_dim', 'feature_dim', '说明'],
        [
            ['default.yaml', 'ViT-B/16', '512', '512', '默认配置，骨干与特征提取器一致'],
            ['vitl14.yaml', 'ViT-L/14', '768', '512', '混合骨干，特征用 ViT-B/16，SVLA 文本用 ViT-L/14'],
        ],
        col_widths=[1.3, 1.2, 1.0, 1.0, 2.0]
    )
    doc.add_paragraph(
        '当 feature_dim ≠ visual_width 时，SVLA.__init__ 自动启用 feat_proj Linear 投影层'
        '将特征维度对齐。validate_clip_config() 校验配置一致性。'
    )

    doc.add_page_break()

    # ==================== 第四章 训练与评估 ====================
    doc.add_heading('第四章 训练与评估', level=1)

    doc.add_heading('4.1 损失函数', level=2)
    doc.add_paragraph(
        'SVLA 训练采用三项损失加权组合，分别负责二分类异常检测、多类别分类与文本特征正则：'
    )
    add_formula(doc, 'loss = CLAS2 + class_loss_alpha · CLASM + txtreg_weight · TextReg')

    doc.add_heading('4.1.1 CLAS2（二分类异常损失）', level=3)
    doc.add_paragraph(
        'CLAS2_dasmil_weighted 实现 π-TopK 加权的 MIL 二分类损失。流程：'
    )
    clas2_steps = [
        '逐镜头 π-TopK 池化：每个镜头内取 top-k 帧分数均值 → 镜头分数',
        '镜头间 π 加权聚合：w = π / π.sum()，video_logit = sum(w · shot_scores)',
        '标签翻转：labels_bin = 1 - labels[:, 0]（normal 标签翻转为异常标签）',
        'Focal Loss：focal_weight = α · (1 - pt)^γ，默认 α=0.25, γ=2.0',
        'loss = BCE(focal_weight · video_logit, labels_bin)',
    ]
    for s in clas2_steps:
        doc.add_paragraph(s, style='List Number')

    doc.add_heading('4.1.2 CLASM（多类别分类损失）', level=3)
    doc.add_paragraph(
        'CLASM_dasmil_weighted 实现逐类别的 π-TopK 加权 MIL 损失。流程：'
    )
    clasm_steps = [
        '逐类别逐镜头 π-TopK 池化：每个类别在每个镜头内取 top-k 帧分数均值',
        '镜头间 π 加权聚合 → 视频级类别 logit',
        '标签归一化：labels = labels / labels.sum(dim=1, keepdim=True)',
        'Focal NLL：loss = -mean(sum(labels · log_softmax(logits)))（可选 Focal 加权）',
    ]
    for s in clasm_steps:
        doc.add_paragraph(s, style='List Number')

    doc.add_heading('4.1.3 TextReg（文本正则）', level=3)
    doc.add_paragraph(
        'text_feature_regularizer 惩罚文本特征间的余弦相似度，包含两项：'
    )
    textreg_points = [
        '原有项：拉远 normal 类与异常类的文本特征（防止 normal 被误判为异常）',
        '新增项：拉远异常类之间的文本特征（防止类别方向坍缩）',
        'loss = (loss_normal + loss_others) · txtreg_weight',
    ]
    for p in textreg_points:
        doc.add_paragraph(p, style='List Bullet')

    doc.add_heading('4.2 优化器与学习率', level=2)
    doc.add_paragraph(
        '采用 AdamW 优化器，分组学习率策略：'
    )
    add_table_from_rows(doc,
        ['参数组', '学习率', '权重衰减', '说明'],
        [
            ['shot_density_head', 'lr × 5.0 (1e-4)', '0', 'π 参数需更快收敛'],
            ['其余 SVLA 参数', 'lr (2e-5)', '0.01', '标准 AdamW 配置'],
            ['CLIP 参数', '冻结', '—', '不参与训练，降低显存'],
        ],
        col_widths=[1.8, 1.5, 1.0, 2.2]
    )
    doc.add_paragraph(
        '学习率调度：MultiStepLR，milestones=[10, 16]，gamma=0.1。'
        '梯度裁剪：max_norm=1.0。早停：patience=10，任一指标（AUC1 或 AUC2）提升即保存 best_model.pth。'
    )

    doc.add_heading('4.3 训练策略', level=2)
    train_strategies = [
        '双 loader 同步迭代：normal_loader 与 anomaly_loader 同步取 batch，确保每个 batch 含正常与异常样本',
        'AMP 禁用：SVLA 与 F.binary_cross_entropy 在 AMP 下冲突，必须 cudnn.amp=false',
        'num_workers=2 + persistent_workers=True：Windows 共享内存限制（error 1455）',
        '梯度裁剪：max_norm=1.0 防止梯度爆炸',
        '早停：patience=10，监控验证集 AUC1 + AUC2',
        '检查点：保存 epoch + model + optimizer + scheduler + config',
    ]
    for s in train_strategies:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading('4.4 评估指标', level=2)
    add_table_from_rows(doc,
        ['指标', '计算方式', '说明'],
        [
            ['AUC1', 'roc_auc_score(labels, logits1)', '二分类异常 ROC 曲线下面积'],
            ['AP1', 'average_precision_score(labels, logits1)', '二分类异常 PR 曲线下面积'],
            ['AUC2', 'roc_auc_score(labels, 1-softmax(logits2)[:,0])', '多类别异常 ROC 曲线下面积'],
            ['AP2', 'average_precision_score(labels, 1-softmax(logits2)[:,0])', '多类别异常 PR 曲线下面积'],
            ['F1', '2·P·R/(P+R)', '按阈值二值化后的 F1'],
            ['Temporal IoU', 'intersection / union', '时序 IoU，评估段定位'],
        ],
        col_widths=[1.2, 3.0, 2.3]
    )
    doc.add_paragraph(
        '评估时 repeat_factor=16 帧级扩展，即每个视频的段级分数重复 16 次对齐到帧级 GT。'
        '评估器导出 results.json + roc_curve.png + pr_curve.png + report.txt。'
    )

    doc.add_page_break()

    # ==================== 第五章 推理管线 ====================
    doc.add_heading('第五章 推理管线', level=1)

    doc.add_heading('5.1 视频预处理', level=2)
    doc.add_paragraph(
        'VideoPreprocessor 负责视频解码、采样与抽帧。支持 decord（GPU 加速）与 OpenCV 回退。'
        '长视频（duration > max_duration=300s）自动降采样：'
    )
    add_formula(doc, 'over_count = (duration - max_duration) / max_duration')
    add_formula(doc, 'divisor = 2 ^ int(over_count)')
    add_formula(doc, 'sample_rate = sample_rate // divisor')
    doc.add_paragraph(
        '采样后通过 uniform_sample 均匀采样到 num_segments（默认 10），生成 timestamps 列表。'
        'API 支持 num_segments 参数覆盖默认值。'
    )

    doc.add_heading('5.2 CLIP 特征提取', level=2)
    doc.add_paragraph(
        'CLIPFeatureExtractor 封装 CLIP 视觉与文本特征提取：'
    )
    fe_points = [
        'extract_visual_features(frames, batch_size=32)：PIL → preprocess → encode_image → L2 归一化',
        'extract_text_features(prompts)：clip.tokenize → encode_token → encode_text(token_emb, tokens) → L2 归一化',
        'compute_similarity(visual, text)：双端 L2 归一化后点积',
        '特征缓存：基于 MD5 hash 的磁盘缓存，避免重复提取',
        'OOM 处理：batch_size 自动减半重试',
    ]
    for p in fe_points:
        doc.add_paragraph(p, style='List Bullet')
    doc.add_paragraph(
        '注意：本地 clip 模块的 encode_text 需双参数 (text=token_embedding, token=原始 tokens)，'
        '与 OpenAI 官方 CLIP 单参数版本不同。'
    )

    doc.add_heading('5.3 SVLA 前向推理', level=2)
    doc.add_paragraph(
        'process_feat 对齐特征到 visual_length=256，送入 SVLA.forward 得到：'
    )
    add_table_from_rows(doc,
        ['输出', '形状', '语义'],
        [
            ['logits1', '[B, T, 1]', '二分类异常 logit（每帧）'],
            ['logits2', '[B, T, 8]', '8 类（含 normal）相似度 logit（每帧）'],
            ['shot_slices', 'List[List[int]]', '镜头边界索引'],
            ['visual_features_norm', '[B, T, D]', '归一化视觉特征（用于 re-rank）'],
        ],
        col_widths=[1.8, 1.5, 3.2]
    )

    doc.add_heading('5.4 异常分数与类别分数计算', level=2)
    add_formula(doc, 'segment_scores = sigmoid(logits1).squeeze(-1)  # [T]')
    add_formula(doc, 'class_scores_raw = softmax(logits2, dim=-1)    # [T, 8]')
    add_formula(doc, 'anomaly_score = max(segment_scores[:valid_length])')
    add_formula(doc, 'P(cat_i | anomaly) = frame_probs[i] / (1 - frame_probs[normal])')
    add_formula(doc, 'category_scores[cat_i] = anomaly_score × P(cat_i | anomaly)')

    doc.add_paragraph(
        '类别分数聚合采用异常分加权 top-k（k=min(3, valid_length)）：'
        'topk_indices = argsort(weighted_valid_scores)[-k:]，'
        'weights = weighted_valid_scores[topk_indices] / sum(weights)，'
        'frame_probs = sum(class_scores_raw[topk_indices] · weights)。'
        '此设计让异常分高的帧主导类别分布。'
    )

    doc.add_heading('5.5 推理增强', level=2)
    doc.add_paragraph('本项目实现 4 项推理增强，无须重新训练，通过配置开关启用：')

    doc.add_heading('5.5.1 分数校准（Isotonic Regression）', level=3)
    doc.add_paragraph(
        'ScoreCalibrator 使用 sklearn IsotonicRegression 将 raw sigmoid 分数映射到真实概率。'
        '离线在 train.csv 上拟合（需 normal + abnormal 两类样本），持久化到 '
        'checkpoints/calibrator.pkl。推理时 calibrated_score = calibrator.transform(anomaly_score)。'
    )
    add_formula(doc, 'calibrated_score = IsotonicRegression.transform(anomaly_score)')

    doc.add_heading('5.5.2 零样本新类别扩展', level=3)
    doc.add_paragraph(
        '通过 CLIP 文本-图像相似度检测训练集外类别（如赌博、毒品）。'
        '在 __init__ 时为每个扩展类别预计算 3-prompt 平均文本特征，'
        '推理时复用 valid_features 计算余弦相似度，top-k 平均后乘 anomaly_score 作为类别分数。'
    )
    add_formula(doc, 'text_feat = mean(extract_text_features(prompts))  # [D]')
    add_formula(doc, 'sim = valid_feats_norm @ text_feat  # [valid_length]')
    add_formula(doc, 'cat_score = anomaly_score · max(0, topk_mean(sim, k=3))')

    doc.add_heading('5.5.3 OOD 检测', level=3)
    doc.add_paragraph(
        '分布外（OOD）检测基于类别分布的熵与最大概率。分布越平均（熵越高、max_prob 越低），'
        '越像分布外内容：'
    )
    add_formula(doc, 'mean_probs = class_scores_raw[:valid_length].mean(axis=0)')
    add_formula(doc, 'ood_score = 0.5 · (1 - max(mean_probs)) + 0.5 · (entropy(mean_probs) / log(num_classes))')
    doc.add_paragraph('默认阈值 0.5，ood_score ≥ threshold 时 is_ood=True。')

    doc.add_heading('5.5.4 关键帧质量加权', level=3)
    doc.add_paragraph(
        '通过 OpenCV 计算每帧的清晰度（Laplacian 方差）、亮度（V 形函数）与信息熵（Shannon），'
        '加权得到帧质量分。低质量帧（模糊/过暗/过曝）在类别聚合时降权：'
    )
    add_formula(doc, 'clarity = min(1, Laplacian_var / 500)')
    add_formula(doc, 'bright_score = max(0, 1 - 2 · |brightness - 0.5|)')
    add_formula(doc, 'entropy_score = min(1, Shannon_entropy / 8)')
    add_formula(doc, 'quality = 0.5 · clarity + 0.3 · bright_score + 0.2 · entropy_score')

    doc.add_paragraph(
        '重要：质量加权仅影响 category_scores 聚合（weighted_valid_scores），'
        '不影响 anomaly_score（基于 raw_valid_scores），避免低质量帧压低异常分导致风险等级降级。'
    )

    doc.add_heading('5.6 CLIP 零样本 re-rank', level=2)
    doc.add_paragraph(
        '为解决 SVLA 的 prompt 偏置问题（如 "aggressive gestures" 与吸烟手部动作相似导致 '
        'Smoke 被误判为 Abusive），本项目实现 CLIP 零样本 re-rank，始终启用：'
    )
    rerank_points = [
        '_init_rerank_text_feats()：为 7 个标准类别预计算 3-prompt 平均文本特征（纯 CLIP 编码）',
        '复用 valid_features（纯 CLIP 视觉特征，512 维）计算余弦相似度',
        'top-k 平均相似度 × anomaly_score 覆盖 category_scores',
        '绕开 SVLA 的 text_prompt_embeddings 偏置，3-prompt 平均稀释 prompt-1 偏置',
        '仅在 is_harmful=True 时执行（正常视频无需 re-rank）',
    ]
    for p in rerank_points:
        doc.add_paragraph(p, style='List Bullet')

    add_code_block(doc,
        '# CLIP 零样本 re-rank 核心逻辑\n'
        'if self.rerank_text_feats and is_harmful:\n'
        '    valid_features = features[:valid_length]\n'
        '    valid_feats_norm = valid_features / (norm(valid_features, axis=-1, keepdims=True) + 1e-10)\n'
        '    rerank_scores = {}\n'
        '    for cat_en, text_feat in self.rerank_text_feats.items():\n'
        '        sim = valid_feats_norm @ text_feat  # [valid_length]\n'
        '        k = min(3, len(sim))\n'
        '        topk_sim = float(np.sort(sim)[-k:].mean())\n'
        '        rerank_scores[cat_en] = anomaly_score * max(0.0, topk_sim)\n'
        '    category_scores = rerank_scores  # 纯覆盖'
    )

    doc.add_heading('5.7 有害段定位与关键帧抽取', level=2)
    doc.add_paragraph(
        '_locate_harmful_segments 采用 3 帧滑动窗平滑 + 阈值检测 + 区域合并：'
    )
    seg_points = [
        '3 帧滑动窗平滑：减少单帧抖动',
        '阈值检测：segment_scores ≥ anomaly_threshold 的帧标记为异常',
        '区域合并：gap < 1s 的相邻异常段合并为一段',
        '段级类别标注：取段内 peak 帧的 class_scores_raw argmax',
        '关键帧抽取：cv2 读取段内 peak 帧保存为 jpg',
    ]
    for p in seg_points:
        doc.add_paragraph(p, style='List Bullet')

    doc.add_heading('5.8 预警报告生成', level=2)
    doc.add_paragraph(
        'AlertGenerator 将 DetectionResult 转换为面向用户的 AlertReport：'
    )
    add_table_from_rows(doc,
        ['字段', '类型', '说明'],
        [
            ['report_id', 'str (UUID4)', '报告唯一 ID'],
            ['video_id', 'str', '视频标识（去扩展名文件名）'],
            ['alert_level', 'str', 'SAFE / LOW / MEDIUM / HIGH'],
            ['anomaly_score', 'float', '异常分数 [0,1]'],
            ['harmful_contents', 'List[Detail]', '有害内容明细'],
            ['summary', 'str', '中文摘要'],
            ['action_suggestion', 'str', '处置建议'],
            ['processing_time', 'float', '处理耗时（秒）'],
        ],
        col_widths=[1.8, 1.5, 3.2]
    )
    doc.add_paragraph('预警等级判定（左闭右开区间）：')
    add_table_from_rows(doc,
        ['等级', '分数范围', '处置建议'],
        [
            ['SAFE', '[0, 0.3)', '内容正常，建议常规监控'],
            ['LOW', '[0.3, 0.5)', '建议标记待审，纳入审核队列等待处理'],
            ['MEDIUM', '[0.5, 0.8)', '建议限制推荐并人工复审，确认内容性质'],
            ['HIGH', '[0.8, 1.0]', '建议立即下架并转人工审核，等待进一步处理'],
        ],
        col_widths=[1.0, 1.5, 4.0]
    )

    doc.add_page_break()

    # ==================== 第六章 API 设计 ====================
    doc.add_heading('第六章 API 设计', level=1)

    doc.add_heading('6.1 RESTful API 概览', level=2)
    doc.add_paragraph('基于 FastAPI 实现，提供 6 个端点，支持视频与图片检测：')
    add_table_from_rows(doc,
        ['方法', '路径', '说明'],
        [
            ['GET', '/api/v1/health', '健康检查（探测服务与模型加载状态）'],
            ['GET', '/api/v1/categories', '获取类别与预警等级'],
            ['POST', '/api/v1/detect', '上传视频检测（支持 threshold + num_segments 覆盖）'],
            ['POST', '/api/v1/detect_image', '上传图片检测（图片视为单帧视频）'],
            ['GET', '/api/v1/reports/{report_id}', '获取预警报告（内存优先，磁盘回退）'],
            ['GET', '/api/v1/keyframes/{filename}', '获取关键帧图片（image/jpeg）'],
        ],
        col_widths=[0.8, 2.8, 3.0]
    )

    doc.add_heading('6.2 核心数据结构', level=2)
    doc.add_paragraph('DetectionResult 包含 14 个字段，含 4 个推理增强字段：')
    add_table_from_rows(doc,
        ['字段', '类型', '说明'],
        [
            ['anomaly_score', 'float', '视频级异常分数 [0,1]'],
            ['is_harmful', 'bool', '是否判定为有害'],
            ['predicted_categories', 'List[str]', '命中阈值的类别英文名（按分数降序）'],
            ['category_scores', 'Dict[str,float]', '7 类条件置信度'],
            ['harmful_segments', 'List[HarmfulSegment]', '超阈值的有害时间段'],
            ['keyframe_paths', 'List[str]', '关键帧图片路径'],
            ['calibrated_score', 'float', '校准后分数（需启用 calibration）'],
            ['ood_score', 'float', 'OOD 分数（需启用 ood）'],
            ['is_ood', 'bool', '是否分布外内容'],
            ['extra_category_info', 'Dict', '零样本扩展类别元信息'],
        ],
        col_widths=[2.0, 2.0, 2.5]
    )

    doc.add_heading('6.3 图片检测 API', level=2)
    doc.add_paragraph(
        'POST /api/v1/detect_image 支持上传图片（jpg/jpeg/png/bmp/webp）执行有害内容检测。'
        '图片视为单帧视频复用 _run_inference 推理管线，返回结构与 /detect 一致。'
    )
    doc.add_paragraph('调用示例：')
    add_code_block(doc,
        'curl -X POST http://localhost:8000/api/v1/detect_image \\\n'
        '    -F "file=@test.jpg" -F "threshold=0.5"'
    )

    doc.add_heading('6.4 错误处理与状态码', level=2)
    add_table_from_rows(doc,
        ['状态码', '说明'],
        [
            ['400', '视频/图片格式不支持或文件无效'],
            ['500', '推理过程内部错误（如显存不足）'],
            ['503', '模型未加载（需 --checkpoint 启动）'],
            ['404', '报告或关键帧不存在'],
        ],
        col_widths=[1.0, 5.5]
    )

    doc.add_page_break()

    # ==================== 第七章 部署与配置 ====================
    doc.add_heading('第七章 部署与配置', level=1)

    doc.add_heading('7.1 环境要求', level=2)
    add_table_from_rows(doc,
        ['组件', '版本', '说明'],
        [
            ['Python', '3.10', '推荐 3.10，最低 3.9'],
            ['PyTorch', '≥ 2.0', '含 torchvision + torchaudio'],
            ['CUDA', '≥ 11.8', '推荐 12.1'],
            ['ffmpeg', '任意', '视频解码依赖'],
            ['GPU 显存', '≥ 8GB', 'ViT-B/16 训练约 6GB，推理约 2GB'],
        ],
        col_widths=[1.5, 1.5, 3.5]
    )

    doc.add_heading('7.2 安装方式', level=2)
    doc.add_paragraph('推荐使用 Conda 安装：')
    add_code_block(doc,
        'conda env create -f environment.yml\n'
        'conda activate nfsw_detector\n'
        '# 或手动创建\n'
        'conda create -n nfsw_detector python=3.10\n'
        'conda activate nfsw_detector\n'
        'conda install pytorch>=2.0 pytorch-cuda=12.1 cudatoolkit>=11.8 ffmpeg -c pytorch -c nvidia\n'
        'pip install -r requirements.txt'
    )
    doc.add_paragraph(
        '注意：项目自带本地 clip/ 模块（含 encode_token + 双参数 encode_text），'
        '无需安装 OpenAI 官方 CLIP。'
    )

    doc.add_heading('7.3 配置文件说明', level=2)
    doc.add_paragraph('configs/default.yaml 包含 12 个配置段：')
    add_table_from_rows(doc,
        ['配置段', '关键参数', '说明'],
        [
            ['model', 'clip_variant, embed_dim, visual_length', 'SVLA 模型架构'],
            ['data', 'dataset, num_segments, supported_formats', '数据集与采样'],
            ['training', 'batch_size, lr, epochs, focal', '训练超参数'],
            ['cuda', 'num_workers, amp, tf32', 'CUDA 与数据加载'],
            ['inference', 'anomaly_threshold, alert_levels', '推理阈值与预警等级'],
            ['demo', 'port, share, max_file_size', 'Gradio Demo'],
            ['logging', 'level, tensorboard', '日志与可视化'],
            ['labels', '7 类中英文映射', '类别定义'],
            ['label_map', '8 键（含 normal）', '类别顺序'],
            ['text_prompts', '7 类各 3 条英文提示', 'CLIP 文本 prompt'],
            ['calibration', 'enabled, path', '分数校准配置'],
            ['ood', 'enabled, threshold', 'OOD 检测配置'],
            ['frame_quality', 'enabled', '帧质量加权配置'],
            ['zero_shot', 'enabled, extra_categories', '零样本扩展配置'],
        ],
        col_widths=[1.3, 2.7, 2.5]
    )

    doc.add_heading('7.4 CLI 命令', level=2)
    add_table_from_rows(doc,
        ['命令', '说明', '关键参数'],
        [
            ['train', '训练模型', '--config, --resume, --device'],
            ['evaluate', '评估模型，输出 AUC/AP', '--config, --checkpoint, --output'],
            ['detect', '单视频检测并生成报告', '--config, --checkpoint, --video, --threshold, --num-segments'],
            ['demo', '启动 Gradio Demo 服务', '--config, --checkpoint, --port, --share'],
            ['serve', '启动 FastAPI RESTful API 服务', '--config, --checkpoint, --host, --port'],
            ['export', '导出 ONNX/TorchScript', '--checkpoint, --format, --output'],
        ],
        col_widths=[1.0, 2.5, 3.0]
    )

    doc.add_page_break()

    # ==================== 第八章 技术创新点总结 ====================
    doc.add_heading('第八章 技术创新点总结', level=1)

    doc.add_heading('8.1 模型层创新（7 项，需训练）', level=2)

    innovations_model = [
        ('Shot-Aware 镜头条件机制',
         '基于余弦相似度的镜头边界检测 + Shot Transformer 段内增强 + ShotDensityHead 镜头密度估计。'
         '将视频的镜头结构信息显式建模，替代传统均匀帧采样。'),
        ('双向 CFA 跨模态适配器',
         'CFATextAdapter（视觉→文本，动态 Prefix + 门控融合）+ TextCFAdapter（文本←视觉，FiLM 调制）。'
         '双向适配使视觉与语言模态深度对齐，优于单向适配。'),
        ('可配置骨干网络',
         '通过 YAML 切换 ViT-B/16 / ViT-L/14，feat_proj 解决维度不匹配，'
         'validate_clip_config 校验配置一致性。支持混合骨干（特征用 B/16，文本用 L/14）。'),
        ('多类别 Focal Loss + 文本 prompt 增强',
         'Focal BCE/NLL（α=0.25, γ=2.0）解决类别不平衡，每类 3 条文本 prompt 用于 CLIP 编码。'),
        ('π-TopK 软 MIL 池化',
         '镜头异常密度 π 自适应调节 top-k 的 k 值。k = k_base · clamp(1+2·(π-0.22), 0.5, 2.0)。'
         '异常密度高的镜头取更多帧，正常镜头取更少帧。'),
        ('Shot-Aware 图卷积邻接',
         '余弦相似度阈值化 + γ·镜头内全连接 + 对称归一化，双分支（shot-aware + 距离邻接）concat。'),
        ('工程鲁棒性细节',
         'CLIP 冻结、分组学习率（π 参数 ×5）、双 loader 同步、梯度裁剪、早停、AMP 禁用、'
         'in-place 禁止、Windows num_workers 限制处理。'),
    ]
    for i, (title, desc) in enumerate(innovations_model, 1):
        doc.add_heading(f'8.1.{i} {title}', level=3)
        doc.add_paragraph(desc)

    doc.add_heading('8.2 推理时增强创新（5 项，无须重训）', level=2)

    innovations_infer = [
        ('分数校准（Isotonic Regression）',
         '离线在 train.csv 拟合 IsotonicRegression，把 raw sigmoid 映射到真实概率。'
         '持久化到 calibrator.pkl，推理时 transform。'),
        ('零样本新类别扩展',
         'CLIP 文本-图像相似度检测训练集外类别（赌博/毒品）。3-prompt 平均文本特征 + top-k 相似度。'),
        ('OOD 检测',
         'ood_score = 0.5·(1-max_prob) + 0.5·norm_entropy。分布越平均越像 OOD。'),
        ('关键帧质量加权',
         'quality = 0.5·clarity + 0.3·brightness + 0.2·entropy。低质量帧降权（仅影响 category 聚合，不降低 anomaly_score）。'),
        ('CLIP 零样本 re-rank',
         '用纯 CLIP 3-prompt 平均文本特征覆盖 logits2 的偏置类别分数。'
         '绕开 SVLA 的 text_prompt_embeddings 偏置，解决 Smoke/Abusive 类别偏置问题。'),
    ]
    for i, (title, desc) in enumerate(innovations_infer, 1):
        doc.add_heading(f'8.2.{i} {title}', level=3)
        doc.add_paragraph(desc)

    doc.add_heading('8.3 关键公式汇总', level=2)
    formulas = [
        ('π-TopK 池化', 'k = round(max(1, int(L/16)+1) · clamp(1+2·(π-0.22), 0.5, 2.0))'),
        ('类别条件概率', 'P(cat_i|anomaly) = frame_probs[i] / (1 - frame_probs[normal])'),
        ('类别分数', 'category_scores[cat_i] = anomaly_score × P(cat_i|anomaly)'),
        ('OOD 分数', 'ood_score = 0.5·(1-max_prob) + 0.5·(entropy/log(num_classes))'),
        ('帧质量', 'quality = 0.5·min(1,lap_var/500) + 0.3·max(0,1-2|bright-0.5|) + 0.2·min(1,entropy/8)'),
        ('总损失', 'loss = CLAS2 + class_loss_alpha·CLASM + txtreg_weight·TextReg'),
    ]
    for name, formula in formulas:
        p = doc.add_paragraph()
        run = p.add_run(f'{name}：')
        run.font.bold = True
        add_formula(doc, formula)

    doc.add_page_break()

    # ==================== 第九章 工程约束与最佳实践 ====================
    doc.add_heading('第九章 工程约束与最佳实践', level=1)
    doc.add_paragraph('本项目在工程实践中总结的 10 条关键约束：')
    constraints = [
        ('cuda.amp 必须为 false', 'SVLA 与 F.binary_cross_entropy 在 AMP 下冲突，会导致数值溢出'),
        ('num_workers ≤ 2 + persistent_workers=True', 'Windows 共享内存限制（error 1455），Linux 可放宽'),
        ('svla.py 禁止 in-place 操作', '_mask_row_normalize 等 in-place 会破坏 autograd 计算图'),
        ('text_list 必须来自 get_prompt_text(label_map)', '推理与训练类别索引必须对齐，否则类别分数错位'),
        ('process_feat 对齐到 visual_length=256', '模型期望固定长度输入，长短视频分别采样/填充'),
        ('CLIP 参数冻结', '训练时仅微调 SVLA 适配层，降低显存并防止 CLIP 退化'),
        ('使用本地 clip 模块，勿装 OpenAI CLIP', '本地版含 encode_token + 双参数 encode_text，API 不同'),
        ('ViT-L/14 需重新训练', '检查点不兼容 ViT-B/16，embed_dim/visual_width 不同'),
        ('文件读取指定 encoding=utf-8', 'YAML 含非 ASCII 字符（中文标签），默认编码会报错'),
        ('raw_valid_scores 与 weighted_valid_scores 分离', '质量加权不降低 anomaly_score，避免风险等级降级'),
    ]
    for i, (title, desc) in enumerate(constraints, 1):
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(f'{title}：')
        run.font.bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ==================== 附录 ====================
    doc.add_heading('附录', level=1)

    doc.add_heading('附录 A 关键超参数表', level=2)
    add_table_from_rows(doc,
        ['超参数', '默认值', '说明'],
        [
            ['visual_length', '256', '模型期望固定帧数'],
            ['embed_dim / visual_width', '512', 'ViT-B/16 维度'],
            ['shot_sim_thresh', '0.90', '镜头边界相似度阈值'],
            ['shot_min_len', '12', '最小镜头长度'],
            ['shot_gamma', '0.05', '镜头邻接权重'],
            ['pi_floor', '0.05', 'π 下限'],
            ['cfa_tau', '0.8', 'CFA 门控温度'],
            ['cfa_beta', '0.8', 'CFA 融合系数'],
            ['cfa_prefix_len', '32', 'CFA Prefix 长度'],
            ['batch_size', '32', '训练批大小'],
            ['lr', '2e-5', '学习率'],
            ['pi_lr_mult', '5.0', 'π 参数学习率倍数'],
            ['txtreg_weight', '0.3', '文本正则权重'],
            ['class_loss_alpha', '1.0', '多类别损失权重'],
            ['focal_alpha / focal_gamma', '0.25 / 2.0', 'Focal Loss 参数'],
            ['scheduler_milestones', '[10, 16]', '学习率衰减节点'],
            ['anomaly_threshold', '0.5', '异常判定阈值'],
            ['alert_levels', '0.3 / 0.5 / 0.8', 'LOW / MEDIUM / HIGH 阈值'],
        ],
        col_widths=[2.2, 1.8, 2.5]
    )

    doc.add_heading('附录 B 类别标签映射表', level=2)
    add_table_from_rows(doc,
        ['ID', '英文 (en)', '中文 (zh)', 'label_map key'],
        [
            ['0', 'Smoke', '吸烟', 'smoke'],
            ['1', 'Blood', '血腥', 'blood'],
            ['2', 'Violent', '暴力', 'violent'],
            ['3', 'Abusive', '辱骂', 'abusive'],
            ['4', 'Sexy', '色情', 'sexy'],
            ['5', 'Money', '金钱诈骗', 'money'],
            ['6', 'Policy', '政治敏感', 'policy'],
            ['—', 'Normal', '正常', 'normal'],
        ],
        col_widths=[0.6, 1.5, 1.5, 2.0]
    )

    doc.add_heading('附录 C 参考文献', level=2)
    refs = [
        'SVLA: Shot-Conditioned Vision-Language Adaptation for Video Anomaly Detection. IJCAI.',
        'Radford A, et al. Learning Transferable Visual Models From Natural Language Supervision (CLIP). ICML 2021.',
        'Sultani W, et al. Real-World Anomaly Detection in Surveillance Videos (DeepMIL). CVPR 2018.',
        'Pavel Iakubovskii. Graph Convolutional Networks for Video Anomaly Detection.',
        'Perez E, et al. FiLM: Visual Reasoning with a General Conditioning Layer. AAAI 2018.',
    ]
    for i, ref in enumerate(refs, 1):
        doc.add_paragraph(f'[{i}] {ref}')

    return doc


def main():
    output_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'docs')
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, 'NFSW_Detector_技术文档.docx')

    print(f'正在生成技术文档：{output_path}')
    doc = build_document()
    doc.save(output_path)
    print(f'✓ 文档已生成：{output_path}')
    print(f'  文件大小：{os.path.getsize(output_path) / 1024:.1f} KB')


if __name__ == '__main__':
    main()
